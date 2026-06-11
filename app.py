import os
import json
import re
import hashlib
import uuid
import shutil
from datetime import date
from flask import Flask, render_template, request, jsonify, send_file
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['RATE_LIMIT_FILE'] = 'rate_limits.json'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
app.config['DAILY_LIMIT'] = 5

# ── API key configuration ────────────────────────────────────────
# For local testing: paste your key here directly.
# For Vercel/production: set the 'gemini_api_key' environment variable — leave this as "".
HARDCODED_API_KEY = ""
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# In-memory session store
session_data = {}

# ────────────────────────────────────────────────────────────────
# Rate limiting  (file-based JSON – no Redis needed)
# ────────────────────────────────────────────────────────────────

def _load_rate_data():
    path = app.config['RATE_LIMIT_FILE']
    if os.path.exists(path):
        try:
            with open(path, 'r') as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def _save_rate_data(data):
    with open(app.config['RATE_LIMIT_FILE'], 'w') as f:
        json.dump(data, f)

def get_client_id():
    """Stable per-client fingerprint from IP + User-Agent."""
    raw = (request.remote_addr or '') + (request.headers.get('User-Agent', ''))
    return hashlib.sha256(raw.encode()).hexdigest()[:32]

def check_rate_limit():
    """Returns (allowed: bool, used: int, limit: int)."""
    client_id = get_client_id()
    today = str(date.today())
    data = _load_rate_data()
    record = data.get(client_id, {})
    if record.get('date') != today:
        record = {'date': today, 'count': 0}
    used = record['count']
    limit = app.config['DAILY_LIMIT']
    return used < limit, used, limit

def increment_rate_limit():
    client_id = get_client_id()
    today = str(date.today())
    data = _load_rate_data()
    record = data.get(client_id, {})
    if record.get('date') != today:
        record = {'date': today, 'count': 0}
    record['count'] += 1
    data[client_id] = record
    _save_rate_data(data)

# ────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────

def extract_text_from_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        reader = PdfReader(filepath)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext in ['.docx', '.doc']:
        doc = Document(filepath)
        return "\n".join(para.text for para in doc.paragraphs)
    return ""

def make_cache_key(resume_text, name, job_role, job_description, preferred_location):
    """SHA-256 of all inputs – same inputs always produce the same key."""
    combined = f"{resume_text}|{name}|{job_role}|{job_description}|{preferred_location}"
    return hashlib.sha256(combined.encode()).hexdigest()

# ────────────────────────────────────────────────────────────────
# DOCX builder – pure Python, no pdflatex required
# ────────────────────────────────────────────────────────────────

def build_optimized_docx(user_info, analysis, resume_text, output_path):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.8)
        section.right_margin  = Inches(0.8)

    # ── Name ──
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(user_info.get('name', 'Candidate'))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # ── Role ──
    role_para = doc.add_paragraph()
    role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = role_para.add_run(user_info.get('job_role', ''))
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0xC0, 0x6C, 0x2C)

    doc.add_paragraph()

    def section_heading(title):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC0, 0x6C, 0x2C)
        p.paragraph_format.space_after = Pt(2)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'E8C99A')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def bullet_item(text):
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(text)
        r.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # ── Summary ──
    section_heading("Professional Summary")
    summary = analysis.get('overall_summary', '')
    for line in summary.split('\n'):
        line = line.strip().lstrip('•-').strip()
        if line:
            bullet_item(line)
    doc.add_paragraph()

    # ── Skills ──
    all_skills = list(dict.fromkeys(
        analysis.get('present_keywords', []) + analysis.get('skills_gap', [])
    ))
    if all_skills:
        section_heading("Key Skills")
        chunks = [all_skills[i:i+4] for i in range(0, len(all_skills), 4)]
        for chunk in chunks:
            p = doc.add_paragraph("   •   ".join(chunk))
            p.runs[0].font.size = Pt(10)
        doc.add_paragraph()

    # ── Experience level ──
    exp = analysis.get('experience_level_match', '')
    if exp:
        section_heading("Experience Level")
        p = doc.add_paragraph(exp)
        p.runs[0].font.size = Pt(10)
        doc.add_paragraph()

    # ── Improvements ──
    improvements = analysis.get('improvements', [])
    if improvements:
        section_heading("Recommended Improvements")
        for item in improvements:
            bullet_item(item)
        doc.add_paragraph()

    # ── Original / Enhanced Resume Content ──
    section_heading("Resume Content")
    for line in resume_text.split('\n'):
        line = line.strip()
        if line:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(10)
            p.paragraph_format.space_after = Pt(1)

    doc.save(output_path)

# ────────────────────────────────────────────────────────────────
# Routes
# ────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    allowed, used, limit = check_rate_limit()
    server_key_available = bool(os.environ.get('GEMINI_API_KEY', '').strip() or HARDCODED_API_KEY.strip())
    return render_template(
        'index.html',
        rate_used=used,
        rate_limit=limit,
        rate_allowed=allowed,
        server_key_available=server_key_available
    )

@app.route('/upload', methods=['POST'])
def upload_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    file = request.files['resume']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
    allowed_ext = {'.pdf', '.docx', '.doc'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        return jsonify({'error': 'Only PDF and DOCX files are allowed'}), 400
    session_id = str(uuid.uuid4())
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{session_id}{ext}")
    file.save(filepath)
    session_data[session_id] = {'resume_path': filepath, 'resume_filename': file.filename}
    return jsonify({'success': True, 'session_id': session_id, 'filename': file.filename})

@app.route('/rate_status', methods=['GET'])
def rate_status():
    allowed, used, limit = check_rate_limit()
    return jsonify({'allowed': allowed, 'used': used, 'limit': limit})

@app.route('/analyze', methods=['POST'])
def analyze():
    data         = request.json
    session_id   = data.get('session_id')
    name         = data.get('name', '').strip()
    job_role     = data.get('job_role', '').strip()
    job_desc     = data.get('job_description', '').strip()
    location     = data.get('preferred_location', '').strip()
    user_api_key = data.get('api_key', '').strip()

    if not session_id or session_id not in session_data:
        return jsonify({'error': 'Invalid session. Please upload your resume first.'}), 400

    # ── API key resolution (priority: user input → env var → hardcoded fallback) ──
    # HARDCODED_API_KEY is checked last so env var always takes precedence in production
    server_key = (
        os.environ.get('GEMINI_API_KEY', '').strip()
        or HARDCODED_API_KEY.strip()
    )
    using_server_key = False

    if user_api_key:
        api_key = user_api_key
    elif server_key:
        allowed, used, limit = check_rate_limit()
        if not allowed:
            return jsonify({
                'error': f'You have used all {limit} free analyses today. Enter your own Gemini API key to continue.',
                'rate_limited': True,
                'used': used,
                'limit': limit
            }), 429
        api_key = server_key
        using_server_key = True
    else:
        return jsonify({
            'error': 'No API key configured. Set the gemini_api_key environment variable or paste your key in the field above.'
        }), 400

    resume_path = session_data[session_id]['resume_path']
    resume_text = extract_text_from_file(resume_path)

    if not resume_text.strip():
        return jsonify({'error': 'Could not extract text from resume. Ensure it is not a scanned image.'}), 400

    # ── Deterministic caching: same inputs → same result ──
    cache_key = make_cache_key(resume_text, name, job_role, job_desc, location)
    if (session_data[session_id].get('cache_key') == cache_key
            and 'analysis' in session_data[session_id]):
        allowed, used, limit = check_rate_limit()
        return jsonify({
            'success': True,
            'analysis': session_data[session_id]['analysis'],
            'cached': True,
            'rate_used': used,
            'rate_limit': limit
        })

    try:
        genai.configure(api_key=api_key)
        # temperature=0 → fully deterministic outputs for the same prompt
        model = genai.GenerativeModel(
            'gemini-2.5-flash',
            generation_config=genai.GenerationConfig(temperature=0)
        )

        prompt = f"""You are a professional ATS resume consultant. Analyse the resume below and return ONLY valid JSON — no markdown, no code fences, no commentary whatsoever.

RESUME:
{resume_text}

CANDIDATE:
- Name: {name}
- Target Role: {job_role}
- Location Preference: {location}
- Job Description: {job_desc}

Return this exact JSON. Use concise bullet-style strings (≤15 words each) for list fields.
{{
  "ats_score": <integer 0-100>,
  "aligned_job_score": <integer 0-100>,
  "missing_keywords": ["keyword1", "keyword2"],
  "present_keywords": ["keyword1", "keyword2"],
  "strengths": [
    "Quantified achievements across key responsibilities",
    "Clear career progression visible in timeline",
    "Relevant tech stack matches target role requirements"
  ],
  "improvements": [
    "Add measurable metrics to every experience bullet",
    "Include missing keyword: system design",
    "Reorder skills — put most relevant tools first",
    "Replace generic objective with a 3-line impact summary",
    "Remove technologies unused for 5+ years"
  ],
  "overall_summary": "• Bullet 1 (max 20 words)\\n• Bullet 2\\n• Bullet 3\\n• Bullet 4\\n• Bullet 5",
  "experience_level_match": "Mid-level — 4 years aligns but team lead experience missing",
  "skills_gap": ["Docker", "Kubernetes", "System Design"],
  "format_issues": ["Inconsistent date formatting", "No LinkedIn URL in header"]
}}"""

        response = model.generate_content(prompt)
        raw = response.text.strip()
        raw = re.sub(r'^```json\s*', '', raw)
        raw = re.sub(r'^```\s*', '', raw)
        raw = re.sub(r'\s*```$', '', raw)

        analysis = json.loads(raw)

        # Increment only on successful server-key usage
        if using_server_key:
            increment_rate_limit()

        session_data[session_id].update({
            'analysis': analysis,
            'resume_text': resume_text,
            'cache_key': cache_key,
            'api_key': api_key,
            'user_info': {
                'name': name,
                'job_role': job_role,
                'job_description': job_desc,
                'preferred_location': location
            }
        })

        allowed, used, limit = check_rate_limit()
        return jsonify({
            'success': True,
            'analysis': analysis,
            'rate_used': used,
            'rate_limit': limit
        })

    except json.JSONDecodeError as e:
        return jsonify({'error': f'AI returned malformed JSON. Try again. Detail: {str(e)}'}), 500
    except Exception as e:
        # Surface the real Gemini error (invalid key, quota, etc.)
        err_msg = str(e)
        if 'API_KEY_INVALID' in err_msg or 'API key not valid' in err_msg:
            return jsonify({'error': 'Invalid Gemini API key. Check your key and try again.'}), 401
        if 'quota' in err_msg.lower() or 'RESOURCE_EXHAUSTED' in err_msg:
            return jsonify({'error': 'Gemini API quota exceeded. Try again later or use a different API key.'}), 429
        return jsonify({'error': f'Analysis failed: {err_msg}'}), 500

@app.route('/generate_resume', methods=['POST'])
def generate_resume():
    data = request.json
    session_id = data.get('session_id')

    if not session_id or session_id not in session_data:
        return jsonify({'error': 'Invalid session'}), 400
    sd = session_data[session_id]
    if 'analysis' not in sd:
        return jsonify({'error': 'Please run analysis first'}), 400

    try:
        analysis   = sd['analysis']
        user_info  = sd['user_info']
        resume_text = sd['resume_text']

        # ── Optionally get AI-enhanced text ──
        enhanced_text = resume_text
        try:
            genai.configure(api_key=sd['api_key'])
            model = genai.GenerativeModel(
                'gemini-2.5-flash',
                generation_config=genai.GenerationConfig(temperature=0)
            )
            enhance_prompt = f"""Rewrite the resume content below, applying the listed improvements.

ORIGINAL RESUME:
{resume_text}

CANDIDATE:
- Name: {user_info['name']}
- Target Role: {user_info['job_role']}

IMPROVEMENTS:
{json.dumps(analysis.get('improvements', []), indent=2)}

KEYWORDS TO ADD NATURALLY:
{json.dumps(analysis.get('missing_keywords', []), indent=2)}

Return plain text only. Use section labels: SUMMARY, SKILLS, EXPERIENCE, EDUCATION, PROJECTS.
No markdown, no JSON. Keep under 700 words."""
            resp = model.generate_content(enhance_prompt)
            enhanced_text = resp.text.strip()
        except Exception:
            pass  # Fall back to original resume text

        safe_name       = re.sub(r'[^a-zA-Z0-9_-]', '_', user_info.get('name', 'resume')) or 'resume'
        output_filename = f"{safe_name}_optimized_resume.docx"
        output_path     = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

        build_optimized_docx(user_info, analysis, enhanced_text, output_path)

        if os.path.exists(output_path):
            sd['output_docx'] = output_path
            return jsonify({'success': True, 'filename': output_filename, 'format': 'docx'})
        return jsonify({'error': 'Failed to create resume document'}), 500

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download(filename):
    safe = re.sub(r'[^a-zA-Z0-9_\-.]', '', filename)
    path = os.path.join(app.config['OUTPUT_FOLDER'], safe)
    if os.path.exists(path):
        return send_file(path, as_attachment=True, download_name=safe)
    return jsonify({'error': 'File not found'}), 404

if __name__ == '__main__':
    app.run(debug=True, port=5000)